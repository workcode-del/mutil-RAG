from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUTPUT = ROOT / "多模态科研论文RAG技术路线简介.docx"

BLUE = "2E74B5"
DARK_BLUE = "1F4D78"
INK = "1F2937"
MUTED = "5B6573"
LIGHT_BLUE = "E8EEF5"
LIGHT_GRAY = "F2F4F7"
WHITE = "FFFFFF"


def set_run_font(run, size=10.5, bold=False, color=INK, italic=False):
    run.font.name = "Calibri"
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), "Calibri")
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), "Calibri")
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    run.font.color.rgb = RGBColor.from_string(color)


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths_dxa):
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths_dxa)))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), "120")
    tbl_ind.set(qn("w:type"), "dxa")
    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths_dxa:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)
    for row in table.rows:
        for index, cell in enumerate(row.cells):
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(widths_dxa[index]))
            tc_w.set(qn("w:type"), "dxa")
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    header = OxmlElement("w:tblHeader")
    header.set(qn("w:val"), "true")
    tr_pr.append(header)


def add_table_text(cell, text, bold=False, color=INK, size=8.6, align=WD_ALIGN_PARAGRAPH.LEFT):
    paragraph = cell.paragraphs[0]
    paragraph.alignment = align
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.paragraph_format.line_spacing = 1.08
    run = paragraph.add_run(text)
    set_run_font(run, size=size, bold=bold, color=color)


def add_hyperlink(paragraph, url, size=8.0, color=BLUE):
    relationship_id = paragraph.part.relate_to(
        url,
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True,
    )
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), relationship_id)
    run = OxmlElement("w:r")
    run_properties = OxmlElement("w:rPr")
    fonts = OxmlElement("w:rFonts")
    fonts.set(qn("w:ascii"), "Calibri")
    fonts.set(qn("w:hAnsi"), "Calibri")
    fonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    run_properties.append(fonts)
    color_node = OxmlElement("w:color")
    color_node.set(qn("w:val"), color)
    run_properties.append(color_node)
    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "single")
    run_properties.append(underline)
    size_node = OxmlElement("w:sz")
    size_node.set(qn("w:val"), str(int(size * 2)))
    run_properties.append(size_node)
    run.append(run_properties)
    text = OxmlElement("w:t")
    text.text = url
    run.append(text)
    hyperlink.append(run)
    paragraph._p.append(hyperlink)


def add_table_links(cell, urls, size=7.6):
    paragraph = cell.paragraphs[0]
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.paragraph_format.line_spacing = 1.0
    for index, url in enumerate(urls):
        if index:
            paragraph.add_run("\n")
        add_hyperlink(paragraph, url, size=size)


def add_body(doc, text, after=6, size=11, bold_lead=None):
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(after)
    paragraph.paragraph_format.line_spacing = 1.25
    if bold_lead and text.startswith(bold_lead):
        lead = paragraph.add_run(bold_lead)
        set_run_font(lead, size=size, bold=True, color=DARK_BLUE)
        rest = paragraph.add_run(text[len(bold_lead):])
        set_run_font(rest, size=size)
    else:
        run = paragraph.add_run(text)
        set_run_font(run, size=size)
    return paragraph


def add_heading(doc, text, level=1):
    paragraph = doc.add_paragraph(style=f"Heading {level}")
    paragraph.paragraph_format.keep_with_next = True
    run = paragraph.add_run(text)
    set_run_font(run, size=16 if level == 1 else 13, bold=True, color=BLUE)
    return paragraph


def add_callout(doc, label, text):
    table = doc.add_table(rows=1, cols=1)
    set_table_geometry(table, [9360])
    cell = table.cell(0, 0)
    set_cell_shading(cell, LIGHT_BLUE)
    paragraph = cell.paragraphs[0]
    paragraph.paragraph_format.space_before = Pt(1)
    paragraph.paragraph_format.space_after = Pt(1)
    paragraph.paragraph_format.line_spacing = 1.1
    lead = paragraph.add_run(f"{label}  ")
    set_run_font(lead, size=10.1, bold=True, color=DARK_BLUE)
    run = paragraph.add_run(text)
    set_run_font(run, size=10.1, color=INK)
    after = doc.add_paragraph()
    after.paragraph_format.space_after = Pt(1)


def add_footer(section):
    footer = section.footer
    paragraph = footer.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)
    run = paragraph.add_run("硕士论文技术路线简介  |  ")
    set_run_font(run, size=8.5, color=MUTED)
    run = paragraph.add_run("第 ")
    set_run_font(run, size=8.5, color=MUTED)
    fld_char1 = OxmlElement("w:fldChar")
    fld_char1.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    fld_char2 = OxmlElement("w:fldChar")
    fld_char2.set(qn("w:fldCharType"), "end")
    run._r.append(fld_char1)
    run._r.append(instr)
    run._r.append(fld_char2)
    tail = paragraph.add_run(" 页")
    set_run_font(tail, size=8.5, color=MUTED)


def build_document():
    doc = Document()
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1.0)
    section.bottom_margin = Inches(1.0)
    section.left_margin = Inches(1.0)
    section.right_margin = Inches(1.0)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)
    add_footer(section)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25
    for name, size, before, after, color in (
        ("Heading 1", 16, 18, 10, BLUE),
        ("Heading 2", 13, 14, 7, BLUE),
    ):
        style = styles[name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)

    title = doc.add_paragraph()
    title.paragraph_format.space_before = Pt(0)
    title.paragraph_format.space_after = Pt(3)
    run = title.add_run("面向科研论文的细粒度多模态 RAG")
    set_run_font(run, size=22, bold=True, color=DARK_BLUE)
    subtitle = doc.add_paragraph()
    subtitle.paragraph_format.space_before = Pt(0)
    subtitle.paragraph_format.space_after = Pt(8)
    run = subtitle.add_run("技术路线、关键算法与创新点简介")
    set_run_font(run, size=12.5, color=MUTED)

    add_callout(
        doc,
        "研究目标",
        "构建可追溯的科研论文多模态检索系统，使问题定位到具体原句、原图、图注和折线图数据，并在严格上下文预算内形成可回答、可引用的跨论文证据集合。",
    )

    add_heading(doc, "1. 整体思路与技术路线", 1)
    add_body(
        doc,
        "系统将科研PDF解析为粗细粒度异构证据图，并把“相关节点召回”和“完整证据选择”拆成两个可独立消融的阶段。主模型采用2025–2026年开源方案，HGT与PCST仅作为轻量基础算子。",
    )
    route = doc.add_paragraph()
    route.alignment = WD_ALIGN_PARAGRAPH.CENTER
    route.paragraph_format.space_before = Pt(3)
    route.paragraph_format.space_after = Pt(6)
    route.paragraph_format.line_spacing = 1.1
    route_text = "MinerU2.5 -> 科研证据图 -> Qwen3-VL召回/原图精排 -> PCST候选 -> 证据闭包/硬预算 -> Qwen3-VL回答"
    run = route.add_run(route_text)
    set_run_font(run, size=9.3, bold=True, color=DARK_BLUE)

    add_heading(doc, "2. 系统模块、算法及论文依据", 1)
    table = doc.add_table(rows=1, cols=4)
    table.style = "Table Grid"
    headers = ["系统环节", "采用算法/模型", "主要作用", "论文或开源依据"]
    for index, text in enumerate(headers):
        cell = table.rows[0].cells[index]
        set_cell_shading(cell, BLUE)
        add_table_text(cell, text, bold=True, color=WHITE, size=8.5, align=WD_ALIGN_PARAGRAPH.CENTER)
    set_repeat_table_header(table.rows[0])
    rows = [
        ("PDF解析与定位", "MinerU2.5 + PyMuPDF", "抽取句子、图片、图注及bbox并回链PDF。", ["https://arxiv.org/abs/2509.22186"]),
        ("折线图结构化", "PP-Chart2Table + 自集成", "重复解析、数值中位数聚合并记录不确定性。", ["https://huggingface.co/docs/transformers/main/model_doc/pp_chart2table", "https://arxiv.org/abs/2605.27298"]),
        ("多模态基础召回", "Qwen3-VL-Embedding-2B", "统一编码文本、原图与查询；Qdrant分类型召回。", ["https://arxiv.org/abs/2601.04720"]),
        ("结构增强索引", "两层HGT + 关系监督", "利用Figure-Caption-Mention-ChartData真实结构边。", ["https://arxiv.org/abs/2003.01332", "https://arxiv.org/abs/2602.04263"]),
        ("原图精排", "Qwen3-VL-Reranker-2B", "直接输入文本、原图或混合证据，使用RRF融合。", ["https://arxiv.org/abs/2601.04720"]),
        ("证据子图选择", "PCST + 闭包 + 硬预算森林", "生成候选骨架，补全必要依赖并跨论文选择。", ["https://aclanthology.org/2025.findings-emnlp.1211/", "https://arxiv.org/abs/2606.15906"]),
        ("回答与引用", "Qwen3-VL-4B/8B或API", "仅基于证据森林回答，并校验evidence_id。", ["https://arxiv.org/abs/2511.21631"]),
    ]
    for row_index, values in enumerate(rows, start=1):
        cells = table.add_row().cells
        if row_index % 2 == 0:
            for cell in cells:
                set_cell_shading(cell, LIGHT_GRAY)
        for index, value in enumerate(values):
            if index == 3:
                add_table_links(cells[index], value)
            else:
                add_table_text(cells[index], value, size=8.15, align=WD_ALIGN_PARAGRAPH.LEFT)
    set_table_geometry(table, [1350, 2040, 2970, 3000])

    doc.add_page_break()

    add_heading(doc, "3. 创新点一：结构监督的科研多模态异构图索引", 1)
    add_body(
        doc,
        "核心方法：Qwen3-VL获得2048维文本/原图基础向量，节点类型投影和两层HGT产生256维结构向量；Figure-Caption、Figure-Mention和Figure-ChartData是真实正关系，同论文错误配对作为难负样本。Qdrant负责全库召回，HGT只增强候选，VL-Reranker再读取原图精排。",
        after=5,
        bold_lead="核心方法：",
    )
    add_body(
        doc,
        "与现有工作的区别：LILaC已提出通用分层组件图与晚交互检索，因此本文不宣称首次细粒度多模态图；贡献限定为科研PDF显式证据关系监督、轻量双空间索引，以及原句、原图和曲线数据级定位。",
        after=7,
        bold_lead="与现有工作的区别：",
    )

    add_heading(doc, "4. 创新点二：证据闭包约束的预算森林检索", 1)
    add_body(
        doc,
        "核心方法：按论文对候选分组，以相关性为节点prize、关系类型为边成本，PCST生成多尺度候选骨架；随后在原始有向图上执行类型化证据闭包。Figure补Caption，ChartData补原图和图注；闭包后重算token与图片成本，再依据相关性、槽位覆盖、实体新颖性和冗余选择跨论文森林。",
        after=5,
        bold_lead="核心方法：",
    )
    add_body(
        doc,
        "与现有工作的区别：MAGE-RAG已研究预算内多模态图导航；本方法不依赖Agent在线动作，强调确定性证据依赖闭包、闭包后重新计费、严格预算与跨论文多答案组合。PCST只是候选工具，不是创新本身。",
        after=7,
        bold_lead="与现有工作的区别：",
    )

    add_heading(doc, "5. 系统实现与实验验证", 1)
    implementation = doc.add_table(rows=4, cols=2)
    implementation.style = "Table Grid"
    implementation_rows = [
        ("数据与解析", "MMDocRAG、Chart-MRAG Bench、PeerQA与SPIQA用于公开验证；私有材料论文集验证跨论文列表题。"),
        ("部署方式", "默认使用一个Python 3.11环境统一运行MinerU、图表解析、Qwen3-VL检索和HGT/PCST；低显存时仅拆分同环境进程。生成器可调用外部API，向量维度为2048。"),
        ("核心对照", "Qwen3-VL flat、VL-Reranker、LILaC简化基线、无/有关系监督HGT；top-k、PPR、PCST、EC-BFR。"),
        ("评价指标", "Sentence/Figure/Joint Recall、Evidence F1、Closure Validity、Slot Coverage、预算违反率、Citation Precision及质量-预算曲线。"),
    ]
    for row_index, (label, value) in enumerate(implementation_rows):
        left, right = implementation.rows[row_index].cells
        set_cell_shading(left, LIGHT_BLUE)
        add_table_text(left, label, bold=True, color=DARK_BLUE, size=8.8, align=WD_ALIGN_PARAGRAPH.CENTER)
        add_table_text(right, value, size=8.55)
    set_table_geometry(implementation, [1700, 7660])

    add_heading(doc, "主要论文依据", 2)
    reference_urls = [
        "https://arxiv.org/abs/2509.22186",
        "https://arxiv.org/abs/2601.04720",
        "https://arxiv.org/abs/2511.21631",
        "https://arxiv.org/abs/2602.04263",
        "https://arxiv.org/abs/2606.15906",
        "https://arxiv.org/abs/2605.27298",
        "https://arxiv.org/abs/2502.14864",
        "https://arxiv.org/abs/2505.16470",
    ]
    reference_paragraph = doc.add_paragraph()
    reference_paragraph.paragraph_format.space_before = Pt(0)
    reference_paragraph.paragraph_format.space_after = Pt(0)
    reference_paragraph.paragraph_format.line_spacing = 1.05
    for index, url in enumerate(reference_urls):
        if index:
            separator = reference_paragraph.add_run("；")
            set_run_font(separator, size=8.3, color=INK)
        add_hyperlink(reference_paragraph, url, size=8.3)

    doc.core_properties.title = "面向科研论文的细粒度多模态RAG技术路线简介"
    doc.core_properties.subject = "硕士论文技术路线、算法依据与创新点"
    doc.core_properties.author = ""
    doc.core_properties.keywords = "多模态RAG, Qwen3-VL, 科研论文检索, 异构证据图, HGT, PCST, 证据闭包"
    doc.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    build_document()
